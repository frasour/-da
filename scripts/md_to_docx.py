# -*- coding: utf-8 -*-
"""
Simple Markdown -> DOCX converter tailored for DESIGN_SPEC.md structure.
Converts headings (#, ##, ###) to Word heading styles and preserves code blocks.
"""
import io
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(ROOT, 'DESIGN_SPEC.md')
OUT_PATH = os.path.join(ROOT, 'DESIGN_SPEC.docx')

if not os.path.exists(MD_PATH):
    print('ERROR: DESIGN_SPEC.md not found')
    raise SystemExit(1)

with io.open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
# ensure Chinese font compatibility
r = style.element.rPr.rFonts
r.set(qn('w:eastAsia'), '微软雅黑')

in_code = False
code_lines = []
for raw in lines:
    line = raw.rstrip('\n')
    if line.strip().startswith('```'):
        in_code = not in_code
        if not in_code:
            # flush code block
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            # set eastAsia font for Consolas fallback
            r = run._r.rPr.rFonts
            r.set(qn('w:eastAsia'), '微软雅黑')
            code_lines = []
        continue
    if in_code:
        code_lines.append(line)
        continue
    if line.startswith('#'):
        # count heading level
        level = 1
        while level <= 6 and line.startswith('#' * level + ' '):
            level += 1
        level -= 1
        text = line[level+1:] if len(line) > level+1 else ''
        # map to Word heading styles
        heading_style = 'Heading {}'.format(min(level, 3))
        doc.add_heading(text, level=min(level, 3))
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        doc.add_paragraph(line)

# Save
doc.save(OUT_PATH)
print('WROTE:', OUT_PATH)
